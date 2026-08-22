#!/usr/bin/env python3
"""
generate_corpus.py — 합성 코퍼스 생성기 v0.1 (2026-08-14)

docs/03 §5 구현:
  1) 각본 dict (문서 63건 × 인사이트 단위 118건, 신호 S1~S6 결정론 배정)
  2) 블록 본문(prose) 확보 — 두 경로, 신호 문장(verbatim)은 항상 이 스크립트가 소유한다
     (a) scripts/corpus_bodies/*.json 에 이미 작성된 본문을 읽는다  ← 기본 경로 (API 키 불필요)
     (b) 없으면 Claude API로 생성 (ANTHROPIC_API_KEY 있을 때만)
     → 어느 경로든 오프셋·재현성이 LLM 출력 품질에 의존하지 않는다
  3) 자가검증: verbatim 정확히 1회 포함 + 블록 구조 + 성별 대명사 금지 (실패 시 명확한 오류)
  4) 산출: backend/data/corpus/*.txt (정본) + *.docx/*.pdf (포장) + manifest.jsonl + ground_truth.jsonl
  5) 포장 렌더 후 텍스트 재추출 → 정본 대조 자가검증

사용 (레포 루트에서):
  python3 scripts/generate_corpus.py --dry-run          # 각본 검증만 (의존성 불필요)
  bash scripts/generate.sh                              # 전체 63건 생성
  bash scripts/generate.sh --limit 3                    # 앞 3개 문서만 스모크 테스트

본문 출처: `scripts/corpus_bodies/*.json` — 형식 {"H01": {"1": "본문…", "2": "본문…"}}.
  이 파일들은 커밋 대상이다 (코퍼스 재생성의 입력이자 재현성의 근거).
API 키(경로 b용): 레포 루트 .env 의 ANTHROPIC_API_KEY (커밋 금지 — docs/06 §5)
주의: 모든 인물·기관은 가상이다. 실명·실기관·실데이터 삽입 금지 (docs/03 §0).
"""

import argparse
import json
import os
import random
import re
import sys
import time
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
import corpus_v3 as V3   # noqa: E402  — v3 각본(인물 300인·문서 320건·조합형 산문)

ROOT = Path(__file__).resolve().parent.parent
CORPUS_DIR = ROOT / "backend" / "data" / "corpus"
BODIES_DIR = ROOT / "scripts" / "corpus_bodies"
MODEL = "claude-sonnet-5"   # 생성 전용. 추출 파이프라인과 동일 모델 계열 (CLAUDE.md 기술 스택)
TEMPERATURE = 0.8           # 원문은 다양해야 하므로 여기만 높게 (docs/03 §5)
PROMPT_VERSION = "corpus-gen@1"  # llm 로깅 규약 대응용 버전 문자열
RNG = random.Random(42)     # 노이즈 주제 선택용 — 시드 고정으로 각본 결정론 유지

# ──────────────────────────────────────────────────────────────────────────────
# 1. 등장인물 (전원 가상)
# ──────────────────────────────────────────────────────────────────────────────

MSLS = V3.MSL_V3   # 12인 (v1의 A1~A5 유지 + A6~A12 신규, 담당 구간 있음)

# ref: (name, specialty, region, setting, institution, city_st)
HCPS = {
    "HCP-001": ("Daniel Whitcomb",  "EPILEPTOLOGY", "NORTHEAST", "ACADEMIC",        "Northbrook Neurology Institute",     "Albany, NY"),
    "HCP-002": ("Grace Lindqvist",  "NEUROLOGY",    "NORTHEAST", "COMMUNITY",       "Riverstone Medical Group",           "Hartford, CT"),
    "HCP-003": ("Omar Haddad",      "EPILEPTOLOGY", "NORTHEAST", "ACADEMIC",        "Beacon Hill Epilepsy Center",        "Providence, RI"),
    "HCP-004": ("Susan Ratliff",    "NEUROLOGY",    "NORTHEAST", "PRIVATE_PRACTICE","Harborlight Neurology",              "Portland, ME"),
    "HCP-005": ("Peter Kowalczyk",  "NEUROLOGY",    "NORTHEAST", "COMMUNITY",       "Green Mountain Neuroscience",        "Burlington, VT"),
    "HCP-006": ("Alice Nguyen-Barr","NEUROLOGY",    "SOUTH",     "PRIVATE_PRACTICE","Old Dominion Neurology Group",       "Richmond, VA"),
    "HCP-007": ("Marcus Boudreaux", "EPILEPTOLOGY", "SOUTH",     "ACADEMIC",        "Bayou Regional Epilepsy Center",     "Baton Rouge, LA"),
    "HCP-008": ("Renee Calloway",   "NEUROLOGY",    "SOUTH",     "COMMUNITY",       "Palmetto Neurology Associates",      "Charleston, SC"),
    "HCP-009": ("Victor Salas",     "EPILEPTOLOGY", "SOUTH",     "ACADEMIC",        "Hill Country Neuroscience Institute","Austin, TX"),
    "HCP-010": ("Imani Okafor",     "NEUROLOGY",    "SOUTH",     "COMMUNITY",       "Smoky Ridge Medical Center",         "Knoxville, TN"),
    "HCP-011": ("Harold Jessup",    "GENERAL",      "SOUTH",     "PRIVATE_PRACTICE","Magnolia Grove Clinic",              "Jackson, MS"),
    "HCP-012": ("Fiona Gallagher",  "EPILEPTOLOGY", "MIDWEST",   "ACADEMIC",        "Great Lakes Epilepsy Institute",     "Madison, WI"),
    "HCP-013": ("Andrei Petrov",    "NEUROLOGY",    "MIDWEST",   "COMMUNITY",       "Prairie Gate Health",                "Des Moines, IA"),
    "HCP-014": ("Naomi Feldman",    "EPILEPTOLOGY", "MIDWEST",   "ACADEMIC",        "Twin Rivers University Hospital",    "Columbus, OH"),
    "HCP-015": ("Caleb Munson",     "NEUROLOGY",    "MIDWEST",   "PRIVATE_PRACTICE","Cedar Ridge Medical Group",          "Springfield, IL"),
    "HCP-016": ("Priya Raghavan",   "NEUROLOGY",    "MIDWEST",   "COMMUNITY",       "North Star Neurology",               "Minneapolis, MN"),
    "HCP-017": ("Walter Brandt",    "GENERAL",      "MIDWEST",   "PRIVATE_PRACTICE","Flint Hills Neuroscience",           "Wichita, KS"),
    "HCP-018": ("Yolanda Reyes",    "NEUROLOGY",    "WEST",      "COMMUNITY",       "Copper Mesa Neurology",              "Mesa, AZ"),
    "HCP-019": ("Stephen Aldrich",  "NEUROLOGY",    "WEST",      "ACADEMIC",        "Cascade Summit Medical Center",      "Tacoma, WA"),
    "HCP-020": ("Mei-Lin Tsang",    "EPILEPTOLOGY", "WEST",      "ACADEMIC",        "Bayview Epilepsy Associates",        "Oakland, CA"),
    "HCP-021": ("Robert Enloe",     "NEUROLOGY",    "WEST",      "PRIVATE_PRACTICE","High Desert Neurology",              "Boise, ID"),
    "HCP-022": ("Hannah Beaufort",  "GENERAL",      "WEST",      "COMMUNITY",       "Silver Peak Clinic",                 "Reno, NV"),
    "HCP-023": ("Jorge Maldonado",  "NEUROLOGY",    "WEST",      "COMMUNITY",       "Redwood Basin Medical Group",        "Sacramento, CA"),
    "HCP-024": ("Katherine Ost",    "EPILEPTOLOGY", "WEST",      "ACADEMIC",        "Sonoran Neuroscience Institute",     "Tucson, AZ"),
}
ROLE_CODE = {"ACADEMIC": "RSM", "COMMUNITY": "RSM", "PRIVATE_PRACTICE": "AR", "GENERAL": "AR"}

HCPS_LEGACY = dict(HCPS)
HCPS = V3.build_roster(HCPS_LEGACY)   # 300인 (001~024 보존 + 025~300 결정론 생성)

NOISE_TOPICS = [
    "upcoming regional epilepsy conference logistics and who is attending",
    "clinic staffing shortage and scheduling backlog",
    "EHR template frustrations when documenting seizure frequency",
    "general reimbursement paperwork burden (no product specifics)",
    "residency teaching duties and journal club plans",
    "patient volume trends this quarter",
    "generic discussion of the ASM landscape without efficacy claims",
    "hospital construction and parking complaints",
    "telehealth follow-up visit logistics",
    "small talk about local weather and travel",
]

# ──────────────────────────────────────────────────────────────────────────────
# 2. 신호 문장 (verbatim — 스크립트가 소유, ground truth의 원천)
#    key: (doc_key, block_index)  → 1문서=1면담 유형은 block_index=1
# ──────────────────────────────────────────────────────────────────────────────

def sig(sid, stype, seg, barrier, text):
    return {"signal_id": sid, "signal_type": stype, "patient_segment": seg, "barrier_type": barrier, "verbatim": text}

S1 = lambda t: sig("S1", "UNMET_NEED", "PEDIATRIC_TRANSITION", None, t)
S2 = lambda t: sig("S2", "SAFETY_CANDIDATE", "UNSPECIFIED", None, t)
S3 = lambda t: sig("S3", "UNMET_NEED", "UNSPECIFIED", None, t)          # post-stroke: v0.1 enum에 없음 → unmapped 유도
S4y = lambda t: sig("S4", "INFO_REQUEST", "PEDIATRIC_TRANSITION", None, t)
S4e = lambda t: sig("S4", "INFO_REQUEST", "ELDERLY_65_PLUS", None, t)
S5 = lambda t: sig("S5", "POSITIVE_OUTCOME", "DRE_2PLUS", None, t)
S6 = lambda t: sig("S6", "TREATMENT_BARRIER", "ELDERLY_65_PLUS", "DDI_CONCERN", t)
X  = lambda sid, t: sig(sid, "CRITIC_BAIT", None, None, t)               # Critic 차단 시연용 의도 삽입 (docs/03 §3.5)

# 대명사 규칙 (2026-08-14): planted 문장과 생성 본문 모두 HCP를 he/she로 지칭하지 않는다.
#   이유 ① 가명 HCP의 성별을 이름으로 추정하지 않는다 ② 각본 문장과 HCP 명부 사이의 성별 불일치 버그를
#   구조적으로 제거한다(실제로 H02·M24 등에서 발생했음). 역할 명사("the physician")·they·무주어를 쓴다.
PLANTED = {
    # ---- S1: 청소년(12–17) 약물난치성 치료 공백 — 14건 / HCP 9인 / 3권역 ----
    ("H01", 2): [S1("Keeps a running list of 15- and 16-year-olds with drug-resistant focal seizures who simply have to wait until 18 before cenobamate is even an option.")],
    ("H02", 2): [S1("Said, and I quote, 'my hardest clinic days are telling parents of a seventeen-year-old that we have nothing new to offer until adulthood.'")],
    ("H03", 1): [S1("Mentioned two adolescent patients who failed carbamazepine and levetiracetam and said the treatment options in that age group feel like a dead end.")],
    ("H05", 4): [S1("Estimated three or four teens in the panel each year hit the same wall: two failed ASMs, still seizing, and no access to cenobamate until age 18.")],
    ("H07", 2): [S1("Said the 12-to-17 group with refractory focal epilepsy is where the biggest gap sits, since the adult-only label leaves them waiting.")],
    ("H09", 3): [S1("Raised, unprompted, that a 16-year-old with drug-resistant focal seizures had to stay on a failing regimen because cenobamate is not approved below 18.")],
    ("H11", 1): [S1("Described the habit of 'parking' adolescent drug-resistant cases until 18, calling it an uncomfortable way to practice.")],
    ("C01", 2): [S1("During Q&A, described a 17-year-old who failed two ASMs and asked, pointedly, when data below 18 would exist for cenobamate.")],
    ("C03", 3): [S1("Standing at the poster, asked why the adolescent drug-resistant population is still excluded, describing families sent away empty-handed until 18."),
                 S4y("Closed by requesting whatever adolescent safety or pharmacokinetic data exists, published or pending.")],
    ("M02", 1): [S1("Called the adolescent transition group the single biggest unmet need — drug-resistant focal seizures at 15 or 16, and cenobamate out of reach until 18.")],
    ("M07", 1): [S1("Sees a steady trickle of drug-resistant teens and called the 18-plus label 'a waiting room with no exit' for them."),
                 S2("Also mentioned one adult patient reporting persistent dizziness after the last titration step, now being watched closely.")],
    ("M15", 1): [S1("Asked me to flag internally that adolescent refractory focal epilepsy keeps coming up with no on-label answer for those families.")],
    ("E03", 1): [S1("Summary of the point raised: repeated adolescent (12-17) drug-resistant cases, nothing to offer on label, families asking why age expansion data does not exist yet.")],
    # ---- S2: 이상사례 시사 — 5건 (S1 단위 2건에 동거: M07, V01) ----
    ("M11", 1): [S2("Brought up an adult who stopped the drug unprompted after a week of feeling foggy and unusually sleepy.")],
    ("P04", 1): [S2("Reported an adult patient complaining of dizziness since the dose increase; advised standard follow-up.")],
    ("H04", 3): [S2("Noted one adult on cenobamate describing daytime somnolence severe enough to affect the daily commute.")],
    # ---- S3: post-stroke epilepsy (v0.1 enum 부재 → SCP 유도) — 6건 / HCP 4인 ----
    ("H02", 4): [S3("Asked whether anyone is looking at post-stroke epilepsy, since more referrals lately are seizures after stroke.")],
    ("H06", 2): [S3("Brought up post-stroke epilepsy again and said these patients do not fit neatly into any of the usual categories.")],
    ("M09", 1): [S3("Described a growing cluster of post-stroke epilepsy patients and asked what evidence exists for cenobamate in that setting.")],
    ("M18", 1): [S3("The interest here is post-stroke seizures — called an overlooked population in every registry reviewed so far.")],
    ("P07", 1): [S3("Quick call, mostly about a post-stroke epilepsy case to discuss at the next visit.")],
    ("E08", 1): [S3("Main theme of the email recap: seizures after stroke are rising in this clinic, with a request for literature on that population.")],
    # ---- S4: 자료 요청 — 8건 (청소년 5: H01.4, C03.3동거, M05, E02, V01 / 노인 DDI 3) ----
    ("H01", 4): [S4y("Asked directly whether any adolescent dosing or safety data for cenobamate is on the horizon.")],
    ("M05", 1): [S4y("Requested a literature package on cenobamate in patients under 18, even early-phase or PK work.")],
    ("E02", 1): [S4y("Action item from the email: send any available adolescent safety data or confirm that none exists yet.")],
    ("M20", 1): [S4e("Asked for a concise interaction table for elderly patients on five or more concomitant medications.")],
    ("P02", 1): [S4e("Wants the DDI summary for geriatric polypharmacy cases before starting anyone new.")],
    ("E06", 1): [S4e("Follow-up requested: interaction and titration guidance for patients over 65 on multiple medications.")],
    # ---- S5: DRE 성인 긍정 신호 — 7건 ----
    ("H03", 4): [S5("Shared that two drug-resistant adults who had failed multiple ASMs are now several months with markedly fewer seizures on cenobamate.")],
    ("H10", 1): [S5("Volunteered that the longest-standing refractory patient in the panel has had the quietest quarter in years since starting cenobamate.")],
    ("M12", 1): [S5("Said several DRE adults in the panel finally responded after years of cycling medications."),
                 X("X2", "Went as far as saying it works for every patient tried on it so far.")],
    ("M22", 1): [S5("Described a drug-resistant adult who went from weekly seizures to rare breakthroughs after titration.")],
    ("P09", 1): [S5("Short call — wanted to share that another refractory adult is responding well on the current dose.")],
    ("C02", 1): [S5("In the hallway, mentioned two treatment-resistant adults doing noticeably better since switching to cenobamate.")],
    # ---- S6: 노인 DDI·적정 부담 (In-label HYP-002) — 10건 / HCP 6인 ----
    ("H01", 1): [S6("The main hesitation with elderly patients is drug-drug interactions on top of already long medication lists.")],
    ("H04", 4): [S6("For older adults, the worry is less about efficacy and more about interaction checks and the burden of slow titration.")],
    ("H06", 4): [S6("Called titration in geriatric patients 'a second job' for caregivers and asked for simpler guidance.")],
    ("H12", 3): [S6("Triple-checks interactions before considering cenobamate in anyone over 65, which slows the decision down.")],
    ("M04", 1): [S6("The hesitation is concomitant medications in elderly patients, with a request for a cleaner way to rule out interactions.")],
    ("M10", 1): [S6("Avoids starting older adults unless a pharmacist reviews the full medication list first, citing DDI concerns.")],
    ("M17", 1): [S6("Elderly polypharmacy came up again, with a request for education materials to hand to caregivers.")],
    ("E01", 1): [S6("Recap: geriatric patients on multiple medications remain the main barrier; requests interaction education for staff.")],
    ("P05", 1): [S6("Brief note — reiterated that interaction burden in patients over 65 keeps this prescriber cautious.")],
    # ---- Critic 차단용 의도 삽입 #1 (허가 범위 벗어난 단정) ----
    ("M24", 1): [X("X1", "The comment was that cenobamate is probably safe for adolescents too, and that stopping the label at 18 seems arbitrary.")],
}

# PII 삽입 (docs/03 §3.4) — 마스킹 정답지 생성 대상. 전부 가상 인물·번호.
PII = {
    "M03": [("NAME", "Marcus Delaney"), ("PHONE", "(555) 214-8890")],
    "E05": [("NAME", "Elena Vargas"), ("PHONE", "(555) 730-4416")],
    "V01": [("NAME", "김도현"), ("PHONE", "010-4132-7789")],
}
PII_SENTENCE = {
    "M03": "The clinic coordinator, Marcus Delaney, left a direct line, (555) 214-8890, for scheduling.",
    "E05": "Please route samples through office manager Elena Vargas, reachable at (555) 730-4416.",
}

# ──────────────────────────────────────────────────────────────────────────────
# 3. 문서 각본 — (doc_key, source_type, date, author, [블록 HCP들])
#    1문서=1면담 유형은 HCP 1명. 렌더 포장: H* 전부 docx, C* 전부 pdf, M01~M08 docx.
# ──────────────────────────────────────────────────────────────────────────────

DOC_PLAN = [
    # 하이라이트 12건 (EN, docx) — 블록 수 [4,5,6,4,5,6,5,5,6,4,6,4] = 60
    ("H01", "HIGHLIGHT_DOC", "2025-09-15", ["A1", "A2"], ["HCP-004", "HCP-001", "HCP-020", "HCP-002"]),
    ("H02", "HIGHLIGHT_DOC", "2025-10-13", ["A2", "A4"], ["HCP-012", "HCP-007", "HCP-018", "HCP-005", "HCP-023"]),
    ("H03", "HIGHLIGHT_DOC", "2025-11-10", ["A1", "A5"], ["HCP-003", "HCP-010", "HCP-020", "HCP-024", "HCP-018", "HCP-006"]),
    ("H04", "HIGHLIGHT_DOC", "2025-12-08", ["A3", "A2"], ["HCP-023", "HCP-012", "HCP-016", "HCP-006"]),
    ("H05", "HIGHLIGHT_DOC", "2026-01-12", ["A4", "A1"], ["HCP-018", "HCP-020", "HCP-024", "HCP-008", "HCP-012"]),
    ("H06", "HIGHLIGHT_DOC", "2026-02-09", ["A5", "A3"], ["HCP-023", "HCP-010", "HCP-021", "HCP-011", "HCP-018", "HCP-020"]),
    ("H07", "HIGHLIGHT_DOC", "2026-05-11", ["A2", "A1"], ["HCP-012", "HCP-013", "HCP-024", "HCP-023", "HCP-006"]),
    ("H08", "HIGHLIGHT_DOC", "2026-03-09", ["A1", "A4"], ["HCP-020", "HCP-009", "HCP-018", "HCP-012", "HCP-010"]),
    ("H09", "HIGHLIGHT_DOC", "2026-06-08", ["A3", "A5"], ["HCP-024", "HCP-023", "HCP-014", "HCP-021", "HCP-016", "HCP-004"]),
    ("H10", "HIGHLIGHT_DOC", "2026-04-13", ["A4", "A2"], ["HCP-020", "HCP-012", "HCP-018", "HCP-005"]),
    ("H11", "HIGHLIGHT_DOC", "2026-07-13", ["A1", "A3"], ["HCP-015", "HCP-010", "HCP-023", "HCP-024", "HCP-016", "HCP-021"]),
    ("H12", "HIGHLIGHT_DOC", "2026-05-25", ["A5", "A4"], ["HCP-018", "HCP-020", "HCP-017", "HCP-012"]),
    # 학회 참관 보고서 3건 (EN, pdf) — 블록 [3,4,3] = 10
    ("C01", "CONGRESS_REPORT", "2025-12-05", ["A1"], ["HCP-012", "HCP-009", "HCP-024"]),
    ("C02", "CONGRESS_REPORT", "2026-04-17", ["A5"], ["HCP-012", "HCP-018", "HCP-023", "HCP-020"]),
    ("C03", "CONGRESS_REPORT", "2026-06-19", ["A3"], ["HCP-016", "HCP-021", "HCP-002"]),
    # 면담 기록 25건 (EN, txt / M01~M08은 docx 포장도)
    ("M01", "MEETING_NOTE", "2025-09-22", ["A1"], ["HCP-019"]),
    ("M02", "MEETING_NOTE", "2025-10-06", ["A2"], ["HCP-001"]),
    ("M03", "MEETING_NOTE", "2025-11-18", ["A3"], ["HCP-024"]),
    ("M04", "MEETING_NOTE", "2025-12-15", ["A4"], ["HCP-019"]),
    ("M05", "MEETING_NOTE", "2026-01-20", ["A5"], ["HCP-003"]),
    ("M06", "MEETING_NOTE", "2025-09-29", ["A2"], ["HCP-022"]),
    ("M07", "MEETING_NOTE", "2026-06-03", ["A1"], ["HCP-007"]),
    ("M08", "MEETING_NOTE", "2025-10-27", ["A4"], ["HCP-011"]),
    ("M09", "MEETING_NOTE", "2026-02-16", ["A3"], ["HCP-016"]),
    ("M10", "MEETING_NOTE", "2026-03-02", ["A5"], ["HCP-022"]),
    ("M11", "MEETING_NOTE", "2026-03-23", ["A2"], ["HCP-023"]),
    ("M12", "MEETING_NOTE", "2026-04-06", ["A1"], ["HCP-012"]),
    ("M13", "MEETING_NOTE", "2025-11-03", ["A4"], ["HCP-017"]),
    ("M14", "MEETING_NOTE", "2025-12-22", ["A3"], ["HCP-005"]),
    ("M15", "MEETING_NOTE", "2026-05-27", ["A5"], ["HCP-013"]),
    ("M16", "MEETING_NOTE", "2026-01-05", ["A2"], ["HCP-008"]),
    ("M17", "MEETING_NOTE", "2026-05-04", ["A1"], ["HCP-004"]),
    ("M18", "MEETING_NOTE", "2026-02-23", ["A3"], ["HCP-021"]),
    ("M19", "MEETING_NOTE", "2026-06-15", ["A4"], ["HCP-002"]),
    ("M20", "MEETING_NOTE", "2026-06-22", ["A5"], ["HCP-006"]),
    ("M21", "MEETING_NOTE", "2026-07-06", ["A2"], ["HCP-014"]),
    ("M22", "MEETING_NOTE", "2026-07-08", ["A1"], ["HCP-018"]),
    ("M23", "MEETING_NOTE", "2026-04-27", ["A4"], ["HCP-009"]),
    ("M24", "MEETING_NOTE", "2026-05-18", ["A3"], ["HCP-012"]),
    ("M25", "MEETING_NOTE", "2026-07-15", ["A5"], ["HCP-015"]),
    # 전화 메모 10건 (EN, txt)
    ("P01", "CALL_NOTE", "2025-10-20", ["A2"], ["HCP-017"]),
    ("P02", "CALL_NOTE", "2026-06-29", ["A4"], ["HCP-011"]),
    ("P03", "CALL_NOTE", "2025-11-24", ["A1"], ["HCP-016"]),
    ("P04", "CALL_NOTE", "2026-03-16", ["A5"], ["HCP-010"]),
    ("P05", "CALL_NOTE", "2026-01-26", ["A3"], ["HCP-011"]),
    ("P06", "CALL_NOTE", "2026-04-20", ["A2"], ["HCP-024"]),
    ("P07", "CALL_NOTE", "2026-02-02", ["A4"], ["HCP-005"]),
    ("P08", "CALL_NOTE", "2026-05-06", ["A1"], ["HCP-020"]),
    ("P09", "CALL_NOTE", "2026-07-01", ["A5"], ["HCP-023"]),
    ("P10", "CALL_NOTE", "2025-12-01", ["A3"], ["HCP-019"]),
    # 이메일 요약 10건 (EN, txt)
    ("E01", "EMAIL_SUMMARY", "2026-01-08", ["A3"], ["HCP-006"]),
    ("E02", "EMAIL_SUMMARY", "2026-05-13", ["A1"], ["HCP-015"]),
    ("E03", "EMAIL_SUMMARY", "2026-07-02", ["A5"], ["HCP-014"]),
    ("E04", "EMAIL_SUMMARY", "2025-10-15", ["A2"], ["HCP-013"]),
    ("E05", "EMAIL_SUMMARY", "2025-11-12", ["A4"], ["HCP-020"]),
    ("E06", "EMAIL_SUMMARY", "2026-06-10", ["A3"], ["HCP-017"]),
    ("E07", "EMAIL_SUMMARY", "2026-03-11", ["A1"], ["HCP-008"]),
    ("E08", "EMAIL_SUMMARY", "2026-02-25", ["A5"], ["HCP-010"]),
    ("E09", "EMAIL_SUMMARY", "2026-07-09", ["A2"], ["HCP-003"]),
    ("E10", "EMAIL_SUMMARY", "2025-12-17", ["A4"], ["HCP-021"]),
    # 음성 전사 3건 (KO, txt, 고정 대본 — API 불필요)
    ("V01", "VOICE_TRANSCRIPT", "2026-07-21", ["A5"], ["HCP-014"]),
    ("V02", "VOICE_TRANSCRIPT", "2026-07-22", ["A5"], ["HCP-019"]),
    ("V03", "VOICE_TRANSCRIPT", "2026-07-23", ["A5"], ["HCP-020"]),
]

# v3 각본 전개 — legacy 63건의 키·날짜·기존 블록·손으로 쓴 본문은 그대로 보존된다.
DOC_PLAN_LEGACY = list(DOC_PLAN)
PLANTED_LEGACY = dict(PLANTED)
_VOICE_SIGS = [("V01", "S1"), ("V01", "S2"), ("V01", "S4"), ("V02", "S6"), ("V03", "S5")]
DOC_PLAN, PLANTED, GEN_BODIES, CONGRESS_META, PLAN_STATS = V3.build_plan(
    HCPS, DOC_PLAN_LEGACY, PLANTED_LEGACY, _VOICE_SIGS)

# docx 포장: 하이라이트 전건(자동) + 단일 면담형 일부. pdf: 학회 전건(자동).
DOCX_ALSO = {f"M0{i}" for i in range(1, 9)} | {
    k for k, st, *_ in DOC_PLAN if st == "MEETING_NOTE" and k.startswith("MX") and int(k[2:]) % 3 == 0}

# 음성 전사 고정 대본 (docs/03 §6) — 한국어, Field 데모 리허설용
VOICE_SCRIPTS = {
    "V01": (
        "MSL: 안녕하세요 선생님, 지난번에 보내드린 자료는 잘 받으셨죠? 오늘도 시간 내주셔서 감사합니다.\n"
        "HCP: 네, 잘 봤어요. 요즘 외래가 많아서 정신이 없네요.\n"
        "MSL: 요즘 난치성 초점발작 청소년 케이스가 좀 있으시다고 들었어요.\n"
        "HCP: 맞아요. 2제 실패한 17세 환자인데 성인 허가라 지금은 손을 쓸 수가 없어요. 18세까지 기다리는 것 말고는 방법이 없네요.\n"
        "MSL: 그런 케이스가 반복되는군요. 기록해 두겠습니다.\n"
        "HCP: 그리고 성인 환자 한 분은 복용 시작하고 나서 어지러움하고 졸림이 꽤 있다고 하셨어요.\n"
        "MSL: 그 부분은 제가 절차대로 안전성 검토 경로로 전달하겠습니다. 자세한 경과는 담당 부서에서 확인드릴 거예요.\n"
        "HCP: 네, 그렇게 해주세요. 아, 그리고 청소년 용량이나 안전성 자료가 나온 게 있으면 꼭 공유 부탁드릴게요.\n"
        "MSL: 확인해서 있는 범위 내에서 정리해 드리겠습니다. 다음 방문은 언제가 좋으실까요?\n"
        "HCP: 다음 주에 김도현 선생님 통해서 잡아주시고, 제 번호 010-4132-7789로 연락 주세요.\n"
        "MSL: 네, 감사합니다 선생님. 조심히 들어가세요.\n"
    ),
    "V02": (
        "MSL: 선생님, 오늘은 고령 환자분들 얘기를 좀 여쭤보고 싶었어요.\n"
        "HCP: 어르신들은 원래 드시는 약이 많아서, 상호작용부터 걱정돼서 시작을 망설이게 된다고 하시더라고요.\n"
        "MSL: 병용약 확인 부담이 크시군요.\n"
        "HCP: 네, 약사 검토까지 거치면 시간이 꽤 걸려요. 간단한 가이드가 있으면 좋겠어요.\n"
        "MSL: 교육 자료 쪽으로 어떤 형태가 필요하신지 다음에 정리해서 가져오겠습니다.\n"
        "HCP: 좋아요. 참, 다음 달 학회는 가세요? 이번엔 세션이 괜찮아 보이던데.\n"
        "MSL: 네, 참석 예정입니다. 학회에서 뵙겠네요.\n"
    ),
    "V03": (
        "MSL: 선생님, 지난번 말씀하신 난치성 환자분은 어떠세요?\n"
        "HCP: 약을 몇 번을 바꿔도 안 잡히던 성인 환자분이 발작 횟수가 눈에 띄게 줄었다고 하셔서 저도 놀랐어요.\n"
        "MSL: 다행이네요. 경과는 계속 지켜보시는 거죠?\n"
        "HCP: 네, 석 달째 유지 중이에요. 물론 환자마다 다르니까 조심스럽긴 하죠.\n"
        "MSL: 네, 개별 경과로 기록해 두겠습니다. 다른 불편사항은 없으셨어요?\n"
        "HCP: 특별한 건 없었어요. 외래 예약 시스템이 바뀌어서 그게 더 골치예요.\n"
        "MSL: 하하, 그건 저도 어쩔 수가 없네요. 다음 방문 때 뵙겠습니다.\n"
    ),
}

# ──────────────────────────────────────────────────────────────────────────────
# 4. 각본 전개 + 검증
# ──────────────────────────────────────────────────────────────────────────────

def month_name(date_str):
    m = ["January","February","March","April","May","June","July","August","September","October","November","December"]
    y, mo, d = date_str.split("-")
    return f"{m[int(mo)-1]} {int(d)}, {y}"

def build_scenario():
    """DOC_PLAN → 문서·블록 단위 각본 리스트. 결정론적."""
    docs = []
    for key, stype, date, authors, hcps in DOC_PLAN:
        blocks = []
        for i, hcp in enumerate(hcps, start=1):
            planted = PLANTED.get((key, i), [])
            blocks.append({
                "block_index": i, "hcp": hcp, "planted": planted,
                "noise_topic": RNG.choice(NOISE_TOPICS),
            })
        fmts = ["TXT"]
        if stype == "HIGHLIGHT_DOC" or key in DOCX_ALSO: fmts.append("DOCX")
        if stype == "CONGRESS_REPORT": fmts.append("PDF")
        docs.append({
            "key": key, "source_type": stype, "date": date, "authors": authors,
            "language": "KO" if stype == "VOICE_TRANSCRIPT" else "EN",
            "blocks": blocks, "formats": fmts,
            "pii": PII.get(key, []),
        })
    return docs

def assert_scenario(docs):
    """docs/03 §2 v3 제약을 결정론적으로 검증. 하나라도 어긋나면 즉시 실패."""
    units = [(d, b) for d in docs for b in d["blocks"]]
    by_key = {d["key"]: d for d in docs}

    def rows(sid):
        out = [(d, b, p) for d, b in units for p in b["planted"] if p["signal_id"] == sid]
        for dk, s2 in [("V01", "S1"), ("V01", "S2"), ("V01", "S4"), ("V02", "S6"), ("V03", "S5")]:
            if s2 == sid and dk in by_key:
                out.append((by_key[dk], by_key[dk]["blocks"][0], None))
        return out

    got = {}
    for sid, (cnt, dhcp, minreg) in V3.TARGETS.items():
        r = rows(sid)
        hcps = {b["hcp"] for _, b, _ in r}
        regs = {HCPS[h][2] for h in hcps}
        got[sid] = (len(r), len(hcps), len(regs))
        assert len(r) == cnt, f"{sid} 건수 {len(r)} != {cnt}"
        if sid in V3.CO_LOCATED:      # 동거 배치로 독립 HCP가 늘어난다 → 하한만 확인
            assert len(hcps) >= dhcp, f"{sid} 독립 HCP {len(hcps)} < {dhcp}"
        else:
            assert len(hcps) == dhcp, f"{sid} 독립 HCP {len(hcps)} != {dhcp}"
        assert len(regs) >= minreg, f"{sid} 권역 {len(regs)} < {minreg}"

    assert len(docs) == 320, f"문서 수 {len(docs)} != 320"
    assert 1000 <= len(units) <= 1200, f"인사이트 단위 {len(units)} 범위 밖"
    # 임계값(반복 ≥5 ∧ 독립 HCP ≥3, UNMET_NEED·TREATMENT_BARRIER)을 넘는 조합이 정확히 4개여야 한다.
    # UNSPECIFIED 환자군은 가설 생성 대상에서 제외된다 (docs/01 §3) — S3(post-stroke)가 여기 해당.
    combos = {}
    for d, b in units:
        for p in b["planted"]:
            if p["signal_type"] not in ("UNMET_NEED", "TREATMENT_BARRIER"):
                continue
            if not p["patient_segment"] or p["patient_segment"] == "UNSPECIFIED":
                continue
            k = (p["patient_segment"], p["signal_type"])
            combos.setdefault(k, {"n": 0, "hcps": set()})
            combos[k]["n"] += 1
            combos[k]["hcps"].add(b["hcp"])
    for dk, sid, seg, st in [("V01", "S1", "PEDIATRIC_TRANSITION", "UNMET_NEED"),
                             ("V02", "S6", "ELDERLY_65_PLUS", "TREATMENT_BARRIER")]:
        k = (seg, st)
        combos.setdefault(k, {"n": 0, "hcps": set()})
        combos[k]["n"] += 1
        combos[k]["hcps"].add(by_key[dk]["blocks"][0]["hcp"])
    crossed = sorted(k for k, v in combos.items() if v["n"] >= 5 and len(v["hcps"]) >= 3)
    expect = sorted([("PEDIATRIC_TRANSITION", "UNMET_NEED"), ("ELDERLY_65_PLUS", "TREATMENT_BARRIER"),
                     ("GENERALIZED_PGTC", "UNMET_NEED"), ("LGS", "UNMET_NEED")])
    assert crossed == expect, f"임계 통과 조합 불일치: {crossed}"

    # 문서 내 HCP 중복 금지 (헤딩 파싱 안정성)
    for d in docs:
        hs = [b["hcp"] for b in d["blocks"]]
        assert len(hs) == len(set(hs)), f"{d['key']}: 문서 내 HCP 중복"
    # 성별 대명사 금지 + verbatim 전역 유일 + 용량 수치 금지
    seen_v = {}
    DOSE = re.compile(r"\b\d+\s?(mg|milligram)", re.IGNORECASE)
    for (dk, bi), plist in PLANTED.items():
        for p in plist:
            m = GENDERED.search(p["verbatim"])
            assert not m, f"PLANTED {dk}.{bi}에 성별 대명사 '{m.group(0) if m else ''}'"
            assert not DOSE.search(p["verbatim"]), f"PLANTED {dk}.{bi}에 용량 수치"
            prev = seen_v.get(p["verbatim"])
            assert prev is None, f"verbatim 중복: {dk}.{bi} == {prev}"
            seen_v[p["verbatim"]] = f"{dk}.{bi}"
    for dk, s2 in PII_SENTENCE.items():
        m = GENDERED.search(s2)
        assert not m, f"PII_SENTENCE {dk}에 성별 대명사"
    # 노이즈 단위가 절반 이상
    sig_u = {(d["key"], b["block_index"]) for d, b in units if b["planted"]}
    sig_u |= {("V01", 1), ("V02", 1), ("V03", 1)}
    noise = len(units) - len(sig_u)
    assert noise >= len(units) // 2, f"노이즈 단위 {noise} < 절반"

    st_count = {}
    for d in docs:
        st_count[d["source_type"]] = st_count.get(d["source_type"], 0) + 1
    return {
        "docs": len(docs), "units": len(units), "hcps": len(HCPS),
        "span": f"{min(d['date'] for d in docs)} ~ {max(d['date'] for d in docs)}",
        "by_type": st_count,
        "signals": {k: {"n": v[0], "hcp": v[1], "regions": v[2]} for k, v in got.items()},
        "hypotheses_crossed": [f"{a}×{b}" for a, b in crossed],
        "noise_units": noise,
        "docx": sum(1 for d in docs if "DOCX" in d["formats"]),
        "pdf": sum(1 for d in docs if "PDF" in d["formats"]),
    }

# ──────────────────────────────────────────────────────────────────────────────
# 5. 텍스트 조립 (스켈레톤은 스크립트 소유 → 헤딩·오프셋 결정론)
# ──────────────────────────────────────────────────────────────────────────────

def heading_line(hcp_ref):
    name, spec, region, setting, inst, city = HCPS[hcp_ref]
    return f"{name}, MD, {inst}, {city} ({ROLE_CODE[setting]})"

def congress_heading(hcp_ref):
    name, _, _, _, inst, _ = HCPS[hcp_ref]
    return f"Conversation — {name}, MD ({inst})"

# 합성 데이터 고지 (2026-08-14) — 문서를 단독으로 열어도 즉시 보이도록 모든 원문의 첫 줄에 넣는다.
# 레포가 public이므로 "실제 기록이 아니다"가 파일 자체에 붙어 있어야 한다.
NOTICE_EN = ("[SYNTHETIC SAMPLE] Fictional record generated for the DELPHi prototype demo. "
             "Names, institutions, and statements are invented; the layout imitates a generic "
             "field-medical format. Not a real document and not derived from any real record.")
NOTICE_KO = ("[합성 샘플] DELPHi 프로토타입 데모용으로 생성된 가상 기록입니다. "
             "인물·기관·발언은 모두 실재하지 않으며, 양식은 일반적인 현장 기록 형식을 흉내 낸 것입니다.")


def doc_header_lines(doc):
    authors = ", ".join(MSLS[a]["name"] for a in doc["authors"])
    date_h = month_name(doc["date"])
    st = doc["source_type"]
    notice = [NOTICE_KO if doc["language"] == "KO" else NOTICE_EN, ""]
    if st == "HIGHLIGHT_DOC":
        return notice + [f"Field Medical Highlights — XCOPRI (cenobamate)",
                         f"Date: {date_h}   |   From: {authors}", "", "FIELD MEDICAL INSIGHTS", ""]
    if st == "CONGRESS_REPORT":
        ab, full, y, mo = CONGRESS_META[doc["key"]]
        mons = ["January","February","March","April","May","June","July","August","September","October","November","December"]
        return notice + [V3.CONGRESS_NOTICE, "",
                         f"Congress Attendance Report — {full} ({ab}) {y}",
                         f"Meeting period: {mons[mo-1]} {y}   |   Author: {authors}", "",
                         "Scope: unsolicited conversations held on-site. Session content and presented data are deliberately not reproduced here.", ""]
    if st == "MEETING_NOTE":
        return notice + [f"Interaction Note   |   Date: {date_h}   |   MSL: {authors}",
                         f"HCP: {heading_line(doc['blocks'][0]['hcp'])}", ""]
    if st == "CALL_NOTE":
        return notice + [f"Call Memo   |   {date_h}   |   {authors}",
                         f"HCP: {heading_line(doc['blocks'][0]['hcp'])}", ""]
    if st == "EMAIL_SUMMARY":
        return notice + [f"Email Summary   |   {date_h}   |   Logged by: {authors}",
                         f"Correspondent: {heading_line(doc['blocks'][0]['hcp'])}", ""]
    if st == "VOICE_TRANSCRIPT":
        return notice + [f"음성 면담 전사   |   {doc['date']}   |   작성: {authors}",
                         f"HCP: {heading_line(doc['blocks'][0]['hcp'])}", "동의 확인: 예 (녹음 전 구두 동의)", ""]
    raise ValueError(st)

def assemble_doc(doc, block_bodies):
    """헤더 + 블록 텍스트를 정본 txt로 조립하고 블록 오프셋을 계산."""
    lines = doc_header_lines(doc)
    text = "\n".join(lines) + "\n"
    spans = []  # (block_index, start, end)
    multi = doc["source_type"] in ("HIGHLIGHT_DOC", "CONGRESS_REPORT")
    for b, body in zip(doc["blocks"], block_bodies):
        start = len(text)
        if multi:
            head = heading_line(b["hcp"]) if doc["source_type"] == "HIGHLIGHT_DOC" else congress_heading(b["hcp"])
            text += head + "\n" + body.rstrip() + "\n\n"
        else:
            text += body.rstrip() + "\n"
        spans.append((b["block_index"], start, len(text)))
    if doc["source_type"] == "CONGRESS_REPORT":
        text += "Follow-ups have been logged in the interaction system.\n"
    return text, spans

# ──────────────────────────────────────────────────────────────────────────────
# 6. LLM 생성 (블록 단위) + 자가검증
# ──────────────────────────────────────────────────────────────────────────────

def load_bodies():
    """scripts/corpus_bodies/*.json 병합 → {doc_key: {block_index(int): body}}"""
    bodies = {}
    if not BODIES_DIR.exists():
        return bodies
    for path in sorted(BODIES_DIR.glob("*.json")):
        data = json.loads(path.read_text(encoding="utf-8"))
        for doc_key, blocks in data.items():
            for bi, body in blocks.items():
                prev = bodies.setdefault(doc_key, {})
                if int(bi) in prev:
                    sys.exit(f"본문 중복 정의: {doc_key} 블록 {bi} (여러 shard에 존재)")
                prev[int(bi)] = body
    return bodies


def load_api_key():
    if os.environ.get("ANTHROPIC_API_KEY"):
        return os.environ["ANTHROPIC_API_KEY"]
    env = ROOT / ".env"
    if env.exists():
        for line in env.read_text().splitlines():
            if line.strip().startswith("ANTHROPIC_API_KEY="):
                return line.split("=", 1)[1].strip().strip('"')
    return None

def block_prompt(doc, block):
    hcp = HCPS[block["hcp"]]
    persona = MSLS[doc["authors"][0]]
    st = doc["source_type"]
    planted = block["planted"]
    plant_lines = "\n".join(f'- "{p["verbatim"]}"' for p in planted)
    pii_line = ""
    if doc["key"] in PII_SENTENCE and block["block_index"] == 1:
        pii_line = f'\nAlso include this sentence VERBATIM once: "{PII_SENTENCE[doc["key"]]}"'
    if st in ("HIGHLIGHT_DOC", "CONGRESS_REPORT"):
        shape = ("Write 2-4 bullet lines, each starting with '> '. No heading line (it is added by the system)."
                 if st == "HIGHLIGHT_DOC" else
                 "Write one short narrative paragraph (2-4 sentences), plain prose, no heading.")
    elif st == "MEETING_NOTE":
        shape = "Write the note body: 120-260 words of free prose, first person MSL voice."
    elif st == "CALL_NOTE":
        shape = "Write 2-4 terse memo lines, each starting with '- '. Fragments are fine."
    else:  # EMAIL_SUMMARY
        shape = "Write a 90-180 word email-style recap paragraph plus a one-line 'Next step:' at the end."
    requirement = (f"These sentences MUST appear VERBATIM, each exactly once, naturally placed:\n{plant_lines}"
                   if planted else
                   f"No product claims here. Topic: {block['noise_topic']}.")
    return f"""You are ghost-writing one fragment of a SYNTHETIC internal field-medical document for a software demo.
Everything is fictional. Do not use real institutions or real people beyond the names given.

Author persona: {persona['name']} — {persona['style']}. Allow a few natural typos (~5%).
Physician context: a {hcp[1].lower().replace('_',' ')} specialist at {hcp[4]} ({hcp[3].lower().replace('_',' ')} setting), {hcp[5]}.
Product context: XCOPRI (cenobamate), adult (18+) focal-onset seizures in the US. Stay at public-label level.
Hard rules: never suggest or encourage off-label use; no efficacy exaggeration beyond any sentence given below; no dosing numbers.
Never refer to the physician as "he" or "she" — use a role noun ("the physician", "this prescriber"), "they", or no subject at all.

{shape}
{requirement}{pii_line}
If a required sentence uses quotes, keep them exactly. Return ONLY the text, no preamble."""

def resolve_body(client, bodies, doc, block):
    gen = GEN_BODIES.get(doc["key"], {}).get(block["block_index"])
    if gen is not None and not bodies.get(doc["key"], {}).get(block["block_index"]):
        validate_block(doc, block, gen)
        return gen
    """본문 확보: 준비된 파일 우선 → 없으면 API. 준비본이 규칙 위반이면 즉시 실패(조용히 API로 넘어가지 않는다)."""
    body = bodies.get(doc["key"], {}).get(block["block_index"])
    if body is not None:
        ok, err = validate_block(doc, block, body)
        if not ok:
            sys.exit(f"본문 검증 실패 — {doc['key']} 블록 {block['block_index']}: {err}\n"
                     f"→ {BODIES_DIR.relative_to(ROOT)} 의 해당 항목을 고치세요.")
        return body
    if client is None:
        sys.exit(f"본문이 없습니다 — {doc['key']} 블록 {block['block_index']}.\n"
                 f"→ {BODIES_DIR.relative_to(ROOT)}/*.json 에 본문을 넣거나, .env에 ANTHROPIC_API_KEY를 설정하세요.")
    return generate_block(client, doc, block)


def generate_block(client, doc, block):
    prompt = block_prompt(doc, block)
    last_err = ""
    for attempt in range(3):
        msg = client.messages.create(
            model=MODEL, max_tokens=700, temperature=TEMPERATURE,
            messages=[{"role": "user", "content": prompt + (f"\n\nPrevious attempt failed: {last_err}. Fix it." if last_err else "")}],
        )
        body = msg.content[0].text.strip()
        ok, last_err = validate_block(doc, block, body)
        if ok:
            return body
        time.sleep(0.5)
    raise RuntimeError(f"{doc['key']}.{block['block_index']} 생성 3회 실패: {last_err}")

GENDERED = re.compile(r"\b(he|she|his|her|hers|him|himself|herself)\b", re.IGNORECASE)

def validate_block(doc, block, body):
    for p in block["planted"]:
        n = body.count(p["verbatim"])
        if n != 1:
            return False, f"verbatim {n}회 포함(정확히 1회 필요): {p['verbatim'][:50]}…"
    if doc["language"] == "EN":
        m = GENDERED.search(body)
        if m:
            return False, f"HCP를 성별 대명사로 지칭함('{m.group(0)}') — 역할 명사·they·무주어를 쓸 것"
    if doc["key"] in PII_SENTENCE and block["block_index"] == 1:
        if body.count(PII_SENTENCE[doc["key"]]) != 1:
            return False, "PII 문장 누락"
    if doc["source_type"] == "HIGHLIGHT_DOC":
        lines = [l for l in body.splitlines() if l.strip()]
        if not all(l.startswith("> ") for l in lines) or not (1 <= len(lines) <= 5):
            return False, "하이라이트 블록은 '> ' 불릿 1~5줄이어야 함"
    if len(body) < 40 or len(body) > 2600:
        return False, f"길이 {len(body)}자 — 범위 밖"
    return True, ""

# ──────────────────────────────────────────────────────────────────────────────
# 7. 포장 렌더 (docx / pdf) + 재추출 검증
# ──────────────────────────────────────────────────────────────────────────────

def render_docx(txt, path):
    from docx import Document
    from docx.shared import Pt
    d = Document()
    for line in txt.splitlines():
        p = d.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(10.5)
        # 볼드 규칙: 제목/섹션 밴드/HCP 헤딩(", MD," 포함) — 실물 양식 근사 (템플릿 정교화는 docs/03 §3.5)
        if line.startswith("Field Medical Highlights") or line == "FIELD MEDICAL INSIGHTS" or ", MD," in line:
            run.font.bold = True
    d.save(path)

def extract_docx(path):
    from docx import Document
    return "\n".join(p.text for p in Document(path).paragraphs)

def render_pdf(txt, path):
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from xml.sax.saxutils import escape
    styles = getSampleStyleSheet()
    story = []
    for line in txt.splitlines():
        if not line.strip():
            story.append(Spacer(1, 8))
        else:
            style = styles["Heading4"] if (line.startswith("Congress Attendance Report") or line.startswith("Conversation —")) else styles["BodyText"]
            story.append(Paragraph(escape(line), style))
    SimpleDocTemplate(str(path), pagesize=letter).build(story)

def extract_pdf(path):
    from pypdf import PdfReader
    return "\n".join((pg.extract_text() or "") for pg in PdfReader(str(path)).pages)

def norm_ws(s):
    return re.sub(r"\s+", " ", s).strip()

def verify_render(doc, txt, corpus_dir, doc_id):
    """포장 파일을 재추출해 정본과 대조. docx=전문 일치 / pdf=신호 문장 포함 (docs/03 §4)."""
    if "DOCX" in doc["formats"]:
        got = extract_docx(corpus_dir / f"{doc_id}.docx")
        assert norm_ws(got) == norm_ws(txt), f"{doc_id}: docx 재추출이 정본과 불일치"
    if "PDF" in doc["formats"]:
        got = norm_ws(extract_pdf(corpus_dir / f"{doc_id}.pdf"))
        for b in doc["blocks"]:
            for p in b["planted"]:
                assert norm_ws(p["verbatim"]) in got, f"{doc_id}: pdf에서 신호 문장 미검출"

# ──────────────────────────────────────────────────────────────────────────────
# 8. 메인 파이프라인
# ──────────────────────────────────────────────────────────────────────────────

def doc_ids(docs):
    """날짜순 전역 시퀀스로 DOC-/INT- id 부여 (결정론)."""
    ordered = sorted(docs, key=lambda d: (d["date"], d["key"]))
    ids = {}
    seq = 0
    for d in ordered:
        seq += 1
        ids[d["key"]] = f"DOC-{d['date'].replace('-', '')}-{seq:03d}"
    return ids

def run(limit=None, dry=False):
    docs = build_scenario()
    stats = assert_scenario(docs)
    print("각본 검증 통과:", json.dumps(stats, ensure_ascii=False))
    if dry:
        print("dry-run 종료 (생성 안 함). 전체 생성: bash scripts/generate.sh")
        return

    bodies = load_bodies()
    n_bodies = sum(len(v) for v in bodies.values())
    print(f"준비된 본문: {n_bodies}개 블록 ({BODIES_DIR.relative_to(ROOT)})")

    client = None
    key = load_api_key()
    if key:
        import anthropic
        client = anthropic.Anthropic(api_key=key)

    CORPUS_DIR.mkdir(parents=True, exist_ok=True)
    ids = doc_ids(docs)
    todo = docs[:limit] if limit else docs
    manifest, ground = [], []
    int_seq = 0

    for di, doc in enumerate(todo, 1):
        doc_id = ids[doc["key"]]
        t0 = time.time()
        if doc["source_type"] == "VOICE_TRANSCRIPT":
            block_bodies = [VOICE_SCRIPTS[doc["key"]] if doc["key"] in VOICE_SCRIPTS
                            else GEN_BODIES[doc["key"]][1]]
        else:
            block_bodies = [resolve_body(client, bodies, doc, b) for b in doc["blocks"]]
        txt, spans = assemble_doc(doc, block_bodies)

        # verbatim 오프셋 (문서 전문 기준 — docs/02 §2) + 블록 범위 검증
        gt_rows = []
        planted_all = [(b, p) for b in doc["blocks"] for p in b["planted"]]
        if doc["key"] == "V01":
            planted_all += [(doc["blocks"][0], s) for s in [
                S1("2제 실패한 17세 환자인데 성인 허가라 지금은 손을 쓸 수가 없어요. 18세까지 기다리는 것 말고는 방법이 없네요."),
                S2("성인 환자 한 분은 복용 시작하고 나서 어지러움하고 졸림이 꽤 있다고 하셨어요."),
                S4y("청소년 용량이나 안전성 자료가 나온 게 있으면 꼭 공유 부탁드릴게요."),
            ]]
        if doc["key"] == "V02":
            planted_all += [(doc["blocks"][0], S6("어르신들은 원래 드시는 약이 많아서, 상호작용부터 걱정돼서 시작을 망설이게 된다고 하시더라고요."))]
        if doc["key"] == "V03":
            planted_all += [(doc["blocks"][0], S5("약을 몇 번을 바꿔도 안 잡히던 성인 환자분이 발작 횟수가 눈에 띄게 줄었다고 하셔서 저도 놀랐어요."))]
        for b, p in planted_all:
            start = txt.find(p["verbatim"])
            assert start != -1 and txt.count(p["verbatim"]) == 1, f"{doc_id}: verbatim 검증 실패"
            span = next(s for s in spans if s[0] == b["block_index"])
            assert span[1] <= start and start + len(p["verbatim"]) <= span[2], f"{doc_id}: verbatim이 블록 범위 밖"
            gt_rows.append({
                "file": f"{doc_id}.txt", "doc_id": doc_id,
                "interaction_ref": f"{doc['key']}#{b['block_index']}",
                "signal_id": p["signal_id"], "verbatim_quote": p["verbatim"],
                "char_start": start, "char_end": start + len(p["verbatim"]),
                "signal_type": p["signal_type"], "patient_segment": p["patient_segment"],
                "barrier_type": p["barrier_type"],
            })

        # PII masked_spans 정답지 (docs/03 §3.4)
        masked = []
        for kind, literal in doc["pii"]:
            i = txt.find(literal)
            assert i != -1, f"{doc_id}: PII 문자열 미검출 ({literal})"
            masked.append({"char_start": i, "char_end": i + len(literal), "kind": kind})

        # 파일 쓰기 + 포장 렌더
        (CORPUS_DIR / f"{doc_id}.txt").write_text(txt, encoding="utf-8")
        if "DOCX" in doc["formats"]:
            render_docx(txt, CORPUS_DIR / f"{doc_id}.docx")
        if "PDF" in doc["formats"]:
            render_pdf(txt, CORPUS_DIR / f"{doc_id}.pdf")
        verify_render(doc, txt, CORPUS_DIR, doc_id)

        # manifest: 한 줄 = 한 interaction (docs/03 §4)
        multi = doc["source_type"] in ("HIGHLIGHT_DOC", "CONGRESS_REPORT")
        for b in doc["blocks"]:
            int_seq += 1
            span = next(s for s in spans if s[0] == b["block_index"])
            block_txt = txt[span[1]:span[2]]
            manifest.append({
                "interaction_id": f"INT-{doc['date'].replace('-','')}-{int_seq:03d}",
                "occurred_on": doc["date"], "hcp_ref": b["hcp"],
                "hcp_specialty": HCPS[b["hcp"]][1], "region": HCPS[b["hcp"]][2],
                "setting": HCPS[b["hcp"]][3], "source_type": doc["source_type"],
                "market": "US", "consent_confirmed": True,
                "language": doc["language"],
                "block_index": b["block_index"] if multi else None,
                "doc_char_start": span[1] if multi else None,
                "doc_char_end": span[2] if multi else None,
                "raw_text": block_txt,
                "masked_spans": [m for m in masked if span[1] <= m["char_start"] < span[2]],
                "file": f"{doc_id}.txt", "source_format": doc["formats"][-1], "doc_id": doc_id,
            })
        ground.extend(gt_rows)
        print(f"[{di}/{len(todo)}] {doc_id} {doc['source_type']:16s} 블록 {len(doc['blocks'])} · {len(gt_rows)}개 신호 · {time.time()-t0:.1f}s")

    with open(CORPUS_DIR / "manifest.jsonl", "w", encoding="utf-8") as f:
        for row in manifest: f.write(json.dumps(row, ensure_ascii=False) + "\n")
    with open(CORPUS_DIR / "ground_truth.jsonl", "w", encoding="utf-8") as f:
        for row in ground: f.write(json.dumps(row, ensure_ascii=False) + "\n")
    print(f"\n완료: 문서 {len(todo)}건 → {CORPUS_DIR}")
    print(f"manifest {len(manifest)}행 · ground_truth {len(ground)}행")
    if limit:
        print("(주의: --limit 실행이라 manifest/ground_truth가 부분본입니다. 전체 재생성 필요)")

if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("--dry-run", action="store_true", help="각본 검증만 (API 불필요)")
    ap.add_argument("--limit", type=int, default=None, help="앞 N개 문서만 생성 (스모크 테스트)")
    a = ap.parse_args()
    run(limit=a.limit, dry=a.dry_run)
